# ***************************************************************************************************
# File        : CommonFunctions.py
# Description : Universal Selenium action helpers with built-in AI locator healing and reporting.
# Author      : Aniket Pathare | Self-Healing AI Framework (2025)
# ***************************************************************************************************

import os
import json
import pandas as pd
from selenium.webdriver.common.by import By
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from selenium.webdriver.chrome.service import Service
from selenium import webdriver
from datetime import datetime
from bs4 import BeautifulSoup
import difflib

# -------------------- GLOBALS --------------------
_LOCATOR_PATH = os.path.join(os.path.dirname(__file__), os.pardir, "Object_Repository", "locators.json")
try:
    with open(_LOCATOR_PATH, "r", encoding="utf-8") as f:
        _LOCATORS = json.load(f)
except Exception:
    _LOCATORS = {}

def get_locator_candidates(key: str):
    """Return all locator candidates for the logical key from repo."""
    if key not in _LOCATORS:
        raise KeyError(f"Locator '{key}' not found in repository.")
    return _LOCATORS[key]

def highlight_element(driver, element, color="red", border_width="3px"):
    try:
        original_style = element.get_attribute("style") or ""
        highlight_style = f"border: {border_width} solid {color};"
        driver.execute_script("arguments[0].setAttribute('style', arguments[1]);", element, highlight_style)
        import time; time.sleep(0.3)
        driver.execute_script("arguments[0].setAttribute('style', arguments[1]);", element, original_style)
    except Exception as e:
        print(f"[DEBUG] Failed to highlight element: {e}")

def get_best_dom_match_by_key(driver, element_key, class_val=None):
    """
    Returns the outerHTML of the element matching the given class_val (exact match)
    and whose attributes or text most closely match the element_key.
    """
    from bs4 import BeautifulSoup
    import difflib

    html = driver.page_source
    soup = BeautifulSoup(html, "html.parser")

    # If class_val is provided, limit candidates to those with exact class match.
    if class_val:
        candidates = []
        for tag in soup.find_all(['button', 'a', 'input']):
            tag_class = tag.get("class")
            if tag_class and " ".join(tag_class) == class_val:
                candidates.append(tag)
        if not candidates:
            # Fallback: nothing found with exact class
            return "<html>ERROR: No element found with exact class match</html>"
    else:
        # fallback if class is not provided, all elements
        candidates = soup.find_all(['button', 'a', 'input'])

    # Now, select the best match by element_key
    best_score = 0
    best_tag = None
    search_key = element_key.lower().replace("_", "").replace("-", "").replace(" ", "")

    for tag in candidates:
        # Check all possible attributes and text for similarity
        attr_string = ""
        for attr in ['id', 'data-qa-id', 'name']:
            attr_string += str(tag.get(attr, ""))
        attr_string += tag.text or ""
        attr_string = attr_string.lower().replace("_", "").replace("-", "").replace(" ", "")

        score = difflib.SequenceMatcher(None, search_key, attr_string).ratio()
        if score > best_score:
            best_score = score
            best_tag = tag

    return str(best_tag) if best_tag else "<html>ERROR: No matching DOM element found</html>"

def extract_dom_snippet_for_healing(driver, locator, max_length=3000, key=None):
    """
    First, extract all DOM elements with exact class match (all classes, same order, space-safe).
    Among those, use the element whose attributes/text best match element_key.
    If no match, fallback to previous logic.
    """
    from selenium.webdriver.common.by import By

    class_val = locator.get("class", "").strip()
    if class_val:
        try:
            # Find ALL elements with exact class match using Selenium
            # XPath for exact class match: @class='class1 class2 ...'
            xpath = f"//*[@class='{class_val}']"
            elems = driver.find_elements(By.XPATH, xpath)
            if elems:
                # If only one element, return immediately
                if len(elems) == 1:
                    html = elems[0].get_attribute("outerHTML")
                    print(f"[HEALING] Only one element found with class '{class_val}'")
                    return html[:max_length]
                # If multiple, use get_best_dom_match_by_key to select best
                # We need the BeautifulSoup representation, so pass HTML source
                # But, limit candidates to elements with exact class only!
                return get_best_dom_match_by_key(driver, key, class_val)[:max_length]
        except Exception as e:
            print(f"[DEBUG] Could not extract DOM by class '{class_val}': {e}")

    # Fallback: old logic, e.g., direct locator
    try:
        by_map = {
            "XPATH": By.XPATH,
            "CSS_SELECTOR": By.CSS_SELECTOR,
            "ID": By.ID,
            "NAME": By.NAME,
            "CLASS_NAME": By.CLASS_NAME,
            "TAG_NAME": By.TAG_NAME,
            "LINK_TEXT": By.LINK_TEXT,
            "PARTIAL_LINK_TEXT": By.PARTIAL_LINK_TEXT,
        }
        by = by_map.get(locator["by"])
        elem = driver.find_element(by, locator["value"])
        snippet = elem.get_attribute("outerHTML")
        if snippet:
            return snippet[:max_length]
        parent = elem.find_element(By.XPATH, "..")
        snippet = parent.get_attribute("outerHTML")
        if snippet:
            return snippet[:max_length]
    except Exception as e:
        print(f"[DEBUG] Could not extract element or parent: {e}")

    # Last fallback: fuzzy key search
    try:
        snippet = get_best_dom_match_by_key(driver, key)
        if snippet:
            print(f"[HEALING] Used fallback fuzzy key match for '{key}'")
            return snippet[:max_length]
    except Exception as e:
        print(f"[DEBUG] Fuzzy key match failed: {e}")

    return "<html>ERROR: could not extract dom</html>"


def smart_find_element(driver, key, timeout=10, test_case_id=None):
    from Self_Healing_AI.healer import ai_only_attempt_heal
    candidates = get_locator_candidates(key)
    by_map = {
        "XPATH": By.XPATH,
        "CSS_SELECTOR": By.CSS_SELECTOR,
        "ID": By.ID,
        "NAME": By.NAME,
        "CLASS_NAME": By.CLASS_NAME,
        "TAG_NAME": By.TAG_NAME,
        "LINK_TEXT": By.LINK_TEXT,
        "PARTIAL_LINK_TEXT": By.PARTIAL_LINK_TEXT,
    }
    for loc in candidates:
        by = by_map.get(loc["by"])
        if not by:
            continue
        try:
            print(f"[DEBUG] Trying '{key}' with ({loc['by']}, '{loc['value']}')")
            elem = WebDriverWait(driver, timeout).until(EC.element_to_be_clickable((by, loc["value"])))
            highlight_element(driver, elem)
            print(f"[DEBUG] Found '{key}' with ({loc['by']}, '{loc['value']}')")
            return elem, loc, False
        except Exception as e:
            print(f"[DEBUG] Failed '{key}' with ({loc['by']}, '{loc['value']}'): {e}")

    # --- AI Healing if all failed ---
    first_orig = candidates[0]
    print(f"[HEALING] All locators failed for '{key}'. Initiating AI healing...")
    snippet = extract_dom_snippet_for_healing(driver, first_orig, max_length=1800, key=key)
    print(f"[HEALING] DOM snippet for '{key}':\n{snippet[:400]} ...\n")
    heal_result = ai_only_attempt_heal(driver, key, (first_orig["by"], first_orig["value"]), snippet, test_case_id)

    if heal_result:
        healed_by, healed_value = heal_result
        by = by_map[healed_by.upper()]
        print(f"[HEALING] Trying AI-suggested locator for '{key}': ({healed_by}, '{healed_value}')")
        try:
            elem = WebDriverWait(driver, timeout + 5).until(EC.element_to_be_clickable((by, healed_value)))
            highlight_element(driver, elem)
            print(f"[HEALING] SUCCESS with AI selector: ({healed_by}, '{healed_value}')")
            healed_loc = {"by": healed_by, "value": healed_value, "auto_healed": True}
            return elem, healed_loc, True
        except Exception as e:
            print(f"[HEALING] FAILED to use AI selector '{healed_value}': {e}")
            driver.save_screenshot(f"healing_fail_{key}.png")
            raise RuntimeError(f"AI healing suggested locator did not work: {e}")
    raise RuntimeError(f"Failed to find element or heal: {key}")


def type_text(driver, key, text, test_case_id=None):
    elem, loc, healed = smart_find_element(driver, key, test_case_id=test_case_id)
    try:
        elem.clear()
    except Exception as e:
        print(f"[DEBUG] Failed to clear field '{key}': {e}")
    elem.send_keys(text)
    return f"Typed in '{key}' using {'AI-healed' if healed else loc['by']} locator."

def click(driver, key, test_case_id=None):
    elem, loc, healed = smart_find_element(driver, key, test_case_id=test_case_id)
    elem.click()
    return f"Clicked '{key}' using {'AI-healed' if healed else loc['by']} locator."

def verify_element(driver, key, test_case_id=None):
    elem, loc, healed = smart_find_element(driver, key, test_case_id=test_case_id)
    return f"Verified '{key}' visible using {'AI-healed' if healed else loc['by']} locator."

def verify_error_method(actual_error_msg: str, expected_error_msg: str):
    actual = actual_error_msg.strip()
    expected = expected_error_msg.strip()
    if expected not in actual:
        raise AssertionError(f"Expected {expected!r} in actual {actual!r}")
    return f"Error message verified: {expected!r}"

# -------------------- REPORTING HELPERS --------------------
def start_word_report(test_case_id: str):
    from docx import Document
    from docx.shared import Inches
    ctrl_df = pd.read_excel(os.path.join("Execution_Control_File", "ExecutionControl.xlsx"))
    ctrl_df.columns = ctrl_df.columns.str.strip()
    desc_row = ctrl_df.loc[ctrl_df["TestCase_ID"].astype(str).str.strip() == test_case_id]
    if desc_row.empty:
        raise KeyError(f"No description found for {test_case_id}")
    description = desc_row.iloc[0]["Description"]
    ts = datetime.now().strftime("%d_%m_%Y_%H_%M_%S")
    docs_dir = os.path.join("Test_Report", "docs")
    os.makedirs(docs_dir, exist_ok=True)
    filename = f"{test_case_id}_{ts}.docx"
    full_path = os.path.join(docs_dir, filename)
    doc = Document()
    doc.add_heading(f"Test Report: {test_case_id}", level=1)
    doc.add_paragraph(f"Description: {description}")
    doc.add_paragraph(f"Started at: {datetime.now():%Y-%m-%d %H:%M:%S}")
    doc.save(full_path)
    return doc, full_path

def add_screenshot_to_report(driver, doc, doc_path: str, test_case_id: str):
    from docx import Document
    from docx.shared import Inches
    ts = datetime.now().strftime("%d_%m_%Y_%H_%M_%S_%f")
    shots_dir = os.path.join("Test_Report", "screenshots")
    os.makedirs(shots_dir, exist_ok=True)
    png_path = os.path.join(shots_dir, f"{test_case_id}_{ts}.png")
    driver.save_screenshot(png_path)
    ctrl_df = pd.read_excel(os.path.join("Execution_Control_File","ExecutionControl.xlsx"))
    ctrl_df.columns = ctrl_df.columns.str.strip()
    desc = ctrl_df.loc[
        ctrl_df["TestCase_ID"].astype(str).str.strip() == test_case_id,
        "Description"
    ].iat[0]
    caption = f"{test_case_id} — {desc}"
    doc.add_paragraph(caption)
    doc.add_picture(png_path, width=Inches(6))
    doc.save(doc_path)

def finalize_word_report(doc, doc_path: str):
    doc.save(doc_path)

# -------------------- BROWSER LAUNCH --------------------
def LaunchBrowser(url: str):
    exe_path = os.getenv('CHROMEDRIVER_PATH', os.path.join(os.getcwd(), 'chromedriver.exe'))
    service = Service(exe_path)
    driver = webdriver.Chrome(service=service)
    driver.maximize_window()
    driver.get(url)
    return driver
